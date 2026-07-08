"""文档解析工具 - 支持PDF、DOCX、TXT格式"""
import io
from typing import Optional
from utils.logger import logger


def parse_pdf(file_content: bytes) -> str:
    """
    解析PDF文件，提取文本内容

    Args:
        file_content: PDF文件二进制内容

    Returns:
        str: 解析后的文本
    """
    try:
        import pdfplumber

        text_parts = []
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        full_text = "\n".join(text_parts)
        logger.debug(f"PDF解析完成: 共{len(pdf.pages)}页, 文本长度{len(full_text)}字")
        return full_text.strip()

    except Exception as e:
        logger.error(f"PDF解析失败: {e}")
        raise ValueError(f"PDF解析失败: {str(e)}")


def parse_docx(file_content: bytes) -> str:
    """
    解析DOCX文件，提取文本内容

    Args:
        file_content: DOCX文件二进制内容

    Returns:
        str: 解析后的文本
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_content))
        text_parts = []

        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        full_text = "\n".join(text_parts)
        logger.debug(f"DOCX解析完成: 共{len(doc.paragraphs)}段, 文本长度{len(full_text)}字")
        return full_text.strip()

    except Exception as e:
        logger.error(f"DOCX解析失败: {e}")
        raise ValueError(f"DOCX解析失败: {str(e)}")


def parse_txt(file_content: bytes) -> str:
    """
    解析TXT文件，提取文本内容

    Args:
        file_content: TXT文件二进制内容

    Returns:
        str: 解析后的文本
    """
    try:
        # 尝试多种编码
        encodings = ["utf-8", "gbk", "gb2312", "utf-16"]
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                logger.debug(f"TXT解析完成(编码:{encoding}): 文本长度{len(text)}字")
                return text.strip()
            except UnicodeDecodeError:
                continue

        raise ValueError("无法识别文本编码")

    except Exception as e:
        logger.error(f"TXT解析失败: {e}")
        raise ValueError(f"TXT解析失败: {str(e)}")


def parse_document(file_content: bytes, file_type: str) -> str:
    """
    根据文件类型解析文档

    Args:
        file_content: 文件二进制内容
        file_type: 文件类型(PDF/DOCX/TXT)

    Returns:
        str: 解析后的文本内容

    Raises:
        ValueError: 不支持的文件类型或解析失败
    """
    file_type = file_type.upper()

    if file_type == "PDF":
        return parse_pdf(file_content)
    elif file_type == "DOCX":
        return parse_docx(file_content)
    elif file_type == "TXT":
        return parse_txt(file_content)
    else:
        raise ValueError(f"不支持的文件类型: {file_type}，仅支持PDF/DOCX/TXT")


def structure_parsed_text(raw_text: str, filename: str) -> str:
    """
    将解析后的原始文本结构化（用于拼接到system message）

    Args:
        raw_text: 原始解析文本
        filename: 文件名

    Returns:
        str: 结构化文本（带来源标识）
    """
    # 简单的结构化：添加来源标识，截取有效内容
    max_length = 2000  # 限制长度避免token过多
    if len(raw_text) > max_length:
        raw_text = raw_text[:max_length] + "...(内容过长已截断)"

    return f"[文档来源: {filename}]\n{raw_text}"